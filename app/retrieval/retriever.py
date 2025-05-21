import os
import logging
import re
import ollama
import pandas as pd
import email
import mimetypes
from io import BytesIO
from typing import List, Dict, Any, Tuple, Optional

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Path to data storage
DATA_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "data")

# Store document content
documents = []
document_titles = []
document_sources = []

# Ollama model to use
OLLAMA_MODEL = "llama3" # Using llama3 which is already available

def extract_text_from_pdf(file_path: str) -> List[Tuple[str, str]]:
    """
    Extract text from PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        List of tuples (section_title, section_content)
    """
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        sections = []
        
        # Extract text from each page
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                title = f"Page {i+1}"
                sections.append((title, text))
                
        # Try to extract document info
        if reader.metadata:
            info = reader.metadata
            title = info.title if info.title else os.path.basename(file_path)
            if title:
                sections.insert(0, ("Document Title", f"{title}"))
        
        return sections
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
        return [("PDF Error", f"Could not extract text from {os.path.basename(file_path)}")]

def extract_text_from_docx(file_path: str) -> List[Tuple[str, str]]:
    """
    Extract text from Word document.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        List of tuples (section_title, section_content)
    """
    try:
        import docx
        
        doc = docx.Document(file_path)
        sections = []
        
        # Try to extract by sections based on headings
        current_heading = "Document Content"
        current_content = []
        
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                # Save previous section
                if current_content:
                    sections.append((current_heading, "\n".join(current_content)))
                    current_content = []
                current_heading = para.text
            else:
                if para.text.strip():
                    current_content.append(para.text)
        
        # Add the last section if it exists
        if current_content:
            sections.append((current_heading, "\n".join(current_content)))
        
        # If no sections were extracted, extract full text
        if not sections:
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            sections.append((os.path.basename(file_path), full_text))
            
        return sections
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        return [("DOCX Error", f"Could not extract text from {os.path.basename(file_path)}")]

def extract_text_from_xlsx(file_path: str) -> List[Tuple[str, str]]:
    """
    Extract data from Excel or CSV file.
    
    Args:
        file_path: Path to the Excel or CSV file
        
    Returns:
        List of tuples (sheet_name, sheet_content_as_text)
    """
    try:
        sections = []
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # For CSV files
        if file_ext == '.csv':
            try:
                # Try to read as CSV
                df = pd.read_csv(file_path, encoding='utf-8', on_bad_lines='warn', low_memory=False)
                logger.info(f"Successfully loaded CSV file {file_path} with shape {df.shape}")
                
                # Generate insights about the CSV
                cols = df.columns.tolist()
                content = [
                    f"CSV File Analysis:",
                    f"Rows: {len(df)}, Columns: {len(cols)}",
                    f"Column Names: {', '.join(cols)}",
                    f"\nData Types:",
                ]
                
                # Add data type information
                for col in cols[:20]:  # Limit to first 20 columns if there are many
                    content.append(f"- {col}: {df[col].dtype}")
                
                # Add statistical summary
                content.append("\nSummary Statistics:")
                try:
                    numeric_cols = df.select_dtypes(include=['number']).columns
                    if not numeric_cols.empty:
                        stats = df[numeric_cols].describe().to_string()
                        content.append(stats)
                except Exception as e:
                    content.append(f"Could not generate statistics: {str(e)}")
                
                # Add sample data (first 10 rows)
                content.append("\nSample Data (first 10 rows):")
                try:
                    sample = df.head(10).to_string(index=True)
                    content.append(sample)
                except Exception as e:
                    content.append(f"Could not display sample: {str(e)}")
                
                sections.append((f"CSV Data", "\n".join(content)))
                
                # Add sections for chunks of data for easier querying
                rows_per_chunk = min(1000, len(df))
                for i in range(0, len(df), rows_per_chunk):
                    chunk = df.iloc[i:i+rows_per_chunk]
                    chunk_desc = f"Rows {i} to {min(i+rows_per_chunk, len(df))}"
                    chunk_content = chunk.head(20).to_string(index=True)  # Only show first 20 rows of each chunk
                    sections.append((f"CSV Chunk: {chunk_desc}", chunk_content))
                
            except Exception as e:
                logger.error(f"Error reading CSV with pandas: {str(e)}")
                # Try alternative approach with Python's csv module
                import csv
                
                with open(file_path, 'r', newline='', encoding='utf-8', errors='replace') as csvfile:
                    try:
                        sample = []
                        reader = csv.reader(csvfile)
                        headers = next(reader, [])
                        
                        # Get first 15 rows as sample
                        for i, row in enumerate(reader):
                            if i < 15:  # Limit to 15 rows for sample
                                sample.append(row)
                            else:
                                break
                        
                        # Build content
                        content = [
                            f"CSV File (basic parsing):",
                            f"Headers: {', '.join(headers)}",
                            "\nSample Data:"
                        ]
                        
                        # Add sample data in formatted way
                        for i, row in enumerate(sample):
                            content.append(f"Row {i+1}: {', '.join(row)}")
                        
                        sections.append((f"CSV Data (Basic Reader)", "\n".join(content)))
                        
                    except Exception as nested_e:
                        logger.error(f"Error with basic CSV reading: {str(nested_e)}")
                        sections.append(("CSV Error", f"Could not parse CSV file {os.path.basename(file_path)}: {str(e)}, {str(nested_e)}"))
        
        # For Excel files
        else:
            try:
                # Read all sheets
                excel_file = pd.ExcelFile(file_path)
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Convert dataframe to text representation
                    content = [
                        f"Sheet: {sheet_name}",
                        f"Rows: {len(df)}, Columns: {len(df.columns)}",
                        f"Column Names: {', '.join(df.columns.tolist())}",
                        "\nData Sample:",
                        df.head(10).to_string(index=True)
                    ]
                    
                    sections.append((f"Excel - {sheet_name}", "\n".join(content)))
            except Exception as e:
                logger.error(f"Error reading Excel file: {str(e)}")
                sections.append(("Excel Error", f"Could not read Excel file {os.path.basename(file_path)}: {str(e)}"))
        
        return sections
    except Exception as e:
        logger.error(f"Error extracting data from {file_path}: {str(e)}")
        return [("Spreadsheet Error", f"Could not extract data from {os.path.basename(file_path)}: {str(e)}")]

def extract_text_from_email(file_path: str) -> List[Tuple[str, str]]:
    """
    Extract content from email file.
    
    Args:
        file_path: Path to the email file
        
    Returns:
        List of tuples (part_name, content)
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            msg = email.message_from_file(f)
        
        sections = []
        
        # Extract header info
        headers = [
            f"From: {msg.get('From', 'Unknown')}",
            f"To: {msg.get('To', 'Unknown')}",
            f"Subject: {msg.get('Subject', 'No Subject')}",
            f"Date: {msg.get('Date', 'Unknown')}"
        ]
        sections.append(("Email Headers", "\n".join(headers)))
        
        # Extract body
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                content_disposition = str(part.get("Content-Disposition"))
                
                # Skip attachments
                if "attachment" in content_disposition:
                    continue
                
                # Get text content
                if content_type == "text/plain":
                    body = part.get_payload(decode=True).decode(errors='ignore')
                    sections.append(("Email Body (Text)", body))
                elif content_type == "text/html":
                    html = part.get_payload(decode=True).decode(errors='ignore')
                    # Could add HTML to text conversion here
                    sections.append(("Email Body (HTML)", html[:1000] + "..."))
        else:
            body = msg.get_payload(decode=True).decode(errors='ignore')
            sections.append(("Email Body", body))
            
        return sections
    except Exception as e:
        logger.error(f"Error extracting content from email {file_path}: {str(e)}")
        return [("Email Error", f"Could not parse email {os.path.basename(file_path)}")]

def extract_text_from_txt(file_path: str) -> List[Tuple[str, str]]:
    """
    Extract text from plaintext file, handling markdown formatting.
    
    Args:
        file_path: Path to the text file
        
    Returns:
        List of tuples (section_title, section_content)
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            content = file.read()
        
        # Check if it's a markdown file with sections
        if re.search(r'#{1,3} ', content):
            # Split into sections by markdown headers
            sections = re.split(r'(#{1,3} .*)', content)
            sections = [s for s in sections if s.strip()]
            
            # Pair headers with content
            result = []
            current_title = os.path.basename(file_path)
            current_content = ""
            
            for section in sections:
                if re.match(r'#{1,3} ', section):
                    if current_content:
                        result.append((current_title, current_content.strip()))
                    current_title = section.lstrip('#').strip()
                    current_content = ""
                else:
                    current_content += section
            
            # Add the last section
            if current_content:
                result.append((current_title, current_content.strip()))
                
            return result
        else:
            # Handle plain text without headers
            paragraphs = [p for p in content.split('\n\n') if p.strip()]
            if len(paragraphs) <= 1:
                # Single paragraph or no clear divisions
                return [(os.path.basename(file_path), content)]
            else:
                # Multiple paragraphs - treat each as a section
                return [(f"Paragraph {i+1}", p) for i, p in enumerate(paragraphs)]
    except Exception as e:
        logger.error(f"Error reading text file {file_path}: {str(e)}")
        return [("Text Error", f"Could not read {os.path.basename(file_path)}")]

def process_file(file_path: str) -> List[Tuple[str, str, str]]:
    """
    Process a file and extract content based on file type.
    
    Args:
        file_path: Path to the file
        
    Returns:
        List of tuples (title, content, source)
    """
    file_name = os.path.basename(file_path)
    file_ext = os.path.splitext(file_name)[1].lower()
    
    # Determine file type and extract content
    sections = []
    source = file_name
    
    try:
        if file_ext == '.pdf':
            sections = extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            sections = extract_text_from_docx(file_path)
        elif file_ext in ['.xlsx', '.xls', '.csv']:
            sections = extract_text_from_xlsx(file_path)
        elif file_ext in ['.eml', '.msg']:
            sections = extract_text_from_email(file_path)
        elif file_ext in ['.txt', '.md', '.markdown']:
            sections = extract_text_from_txt(file_path)
        else:
            # Try to guess based on content
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type:
                if 'text' in mime_type:
                    sections = extract_text_from_txt(file_path)
                elif 'pdf' in mime_type:
                    sections = extract_text_from_pdf(file_path)
                else:
                    sections = [("Unknown File", f"File {file_name} has unsupported format: {mime_type}")]
            else:
                sections = [("Unknown File", f"File {file_name} has an unsupported format")]
    except Exception as e:
        logger.error(f"Error processing file {file_path}: {str(e)}")
        sections = [("Processing Error", f"Error processing {file_name}: {str(e)}")]
    
    # Include source information with each section
    return [(title, content, source) for title, content in sections]

def initialize_index():
    """Initialize by loading documents from the data directory."""
    global documents, document_titles, document_sources
    
    # Clear existing data
    documents.clear()
    document_titles.clear()
    document_sources.clear()
    
    try:
        # Load documents from the data directory
        logger.info(f"Loading documents from {DATA_PATH}")
        
        if not os.path.exists(DATA_PATH):
            logger.warning(f"Data directory {DATA_PATH} does not exist. Creating it.")
            os.makedirs(DATA_PATH, exist_ok=True)
        
        # Process all files in the directory
        file_count = 0
        section_count = 0
        
        for item in os.listdir(DATA_PATH):
            item_path = os.path.join(DATA_PATH, item)
            
            if os.path.isfile(item_path):
                # Process the file
                sections = process_file(item_path)
                for title, content, source in sections:
                    documents.append(content)
                    document_titles.append(title)
                    document_sources.append(source)
                    section_count += 1
                file_count += 1
            elif os.path.isdir(item_path):
                # Recursively process subdirectories
                for root, _, files in os.walk(item_path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        sections = process_file(file_path)
                        for title, content, source in sections:
                            documents.append(content)
                            document_titles.append(title)
                            document_sources.append(source)
                            section_count += 1
                        file_count += 1
        
        logger.info(f"Loaded {section_count} document sections from {file_count} files in {DATA_PATH}")
        
        # Check if Ollama is available
        try:
            models = ollama.list()
            logger.info(f"Ollama is available, using model: {OLLAMA_MODEL}")
            logger.info(f"Available models: {', '.join([model['name'] for model in models['models']])}")
        except Exception as e:
            logger.warning(f"Ollama is not available: {str(e)}. Falling back to keyword matching.")
    
    except Exception as e:
        logger.error(f"Error loading documents: {str(e)}")
        # Create a fallback document if loading fails
        documents = [
            "Artificial Intelligence (AI) refers to the simulation of human intelligence in machines.",
            "Machine Learning is a specific subset of AI that focuses on the development of algorithms and statistical models that enable computers to perform specific tasks without using explicit instructions, relying on patterns and inference instead.",
            "Deep Learning is a specialized subset of machine learning that uses neural networks with many layers.",
            "Natural Language Processing (NLP) is a field of AI focusing on interaction between computers and human language.",
            "Computer Vision involves training computers to interpret and understand the visual world."
        ]
        document_titles = [
            "Artificial Intelligence (AI)",
            "Machine Learning (ML)",
            "Deep Learning",
            "Natural Language Processing (NLP)",
            "Computer Vision"
        ]
        document_sources = ["fallback"] * 5

def retrieve_relevant_context(question: str, max_results: int = 3) -> List[Dict[str, str]]:
    """
    Retrieve the most relevant content based on keyword matching.
    
    Args:
        question: The question to find content for
        max_results: Maximum number of results to return
        
    Returns:
        List of dicts with title, content and source
    """
    if not documents:
        return []
    
    # Simple keyword matching
    keywords = re.findall(r'\b\w+\b', question.lower())
    # Filter out common words and short words
    keywords = [k for k in keywords if len(k) > 2 and k not in [
        "what", "when", "where", "which", "how", "who", "why",
        "this", "that", "with", "from", "will", "would", "should", 
        "could", "about", "the", "and", "for", "are", "you", "your"
    ]]
    
    if not keywords:
        return []
        
    logger.info(f"Extracted keywords: {keywords}")
    
    # Score paragraphs by keyword matches
    scores = []
    for i, doc in enumerate(documents):
        # Count keywords in document
        keyword_count = sum(1 for keyword in keywords if keyword.lower() in doc.lower())
        # Also check title for keywords
        title_count = sum(2 for keyword in keywords if i < len(document_titles) and keyword.lower() in document_titles[i].lower())
        # Combined score with title matches weighted higher
        score = keyword_count + title_count
        scores.append((score, i))
    
    # Sort by score (descending)
    scores.sort(reverse=True)
    
    # Get top relevant paragraphs
    relevant_contexts = []
    
    for score, idx in scores[:max_results]:
        if score > 0:
            title = document_titles[idx] if idx < len(document_titles) else "Section"
            content = documents[idx]
            source = document_sources[idx] if idx < len(document_sources) else "Unknown"
            relevant_contexts.append({
                "title": title,
                "content": content,
                "source": source
            })
    
    return relevant_contexts

def get_answer_from_ollama(question: str, contexts: List[Dict[str, str]]) -> str:
    """
    Use Ollama to generate an answer based on the retrieved contexts.
    
    Args:
        question: The user's question
        contexts: List of relevant document contexts
        
    Returns:
        Generated answer
    """
    try:
        # Format the context and question for the prompt
        formatted_context = ""
        for ctx in contexts:
            formatted_context += f"## {ctx['title']} (from {ctx['source']})\n{ctx['content']}\n\n"
        
        prompt = f"""Based on the following information, please answer the question accurately.
        
Information:
{formatted_context}

Question: {question}

Answer:"""
        
        # Call Ollama API
        response = ollama.generate(
            model=OLLAMA_MODEL,
            prompt=prompt,
            options={
                "temperature": 0.2,  # Low temperature for more factual responses
                "top_k": 50,
                "top_p": 0.9,
                "num_predict": 512  # Limit response length
            }
        )
        
        answer = response['response'].strip()
        return answer
    
    except Exception as e:
        logger.error(f"Error generating answer with Ollama: {str(e)}")
        return None

def get_answer(question: str) -> str:
    """
    Retrieve and answer a question using Ollama and relevant document contexts.
    
    Args:
        question: The question to answer
        
    Returns:
        A string containing the answer
    """
    if not documents:
        return f"No documents loaded. This is a mock answer for: {question}"
    
    try:
        # Retrieve relevant contexts
        relevant_contexts = retrieve_relevant_context(question, max_results=3)
        
        if not relevant_contexts:
            return f"I couldn't find specific information about that in our knowledge base. Please try rephrasing your question."
        
        # Try to use Ollama for the answer
        ollama_answer = get_answer_from_ollama(question, relevant_contexts)
        
        if ollama_answer:
            return ollama_answer
        
        # Fallback to simple context concatenation if Ollama fails
        response = "Based on the available information:\n\n"
        for ctx in relevant_contexts:
            response += f"**{ctx['title']}** (from {ctx['source']})\n{ctx['content']}\n\n"
        
        return response.strip()
    
    except Exception as e:
        logger.error(f"Error generating answer: {str(e)}")
        return f"Sorry, I encountered an error while trying to answer: {question}"

# Add a function to reload the index
def reload_index() -> Dict[str, Any]:
    """
    Reload the document index to pick up new files.
    
    Returns:
        Status information about the reloading process
    """
    try:
        old_count = len(documents)
        initialize_index()
        new_count = len(documents)
        return {
            "success": True,
            "message": f"Successfully reloaded index. Documents: {old_count} → {new_count}",
            "document_count": new_count
        }
    except Exception as e:
        logger.error(f"Error reloading index: {str(e)}")
        return {
            "success": False,
            "message": f"Error reloading index: {str(e)}",
            "document_count": len(documents)
        }

# Initialize index on module import
initialize_index()